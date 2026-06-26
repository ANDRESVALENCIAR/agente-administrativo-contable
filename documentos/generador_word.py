"""
Generador de documentos Word con python-docx.
"""
import logging
import os

from docx import Document

logger = logging.getLogger(__name__)


def generar_word_texto(path: str, titulo: str, contenido: str) -> str:
    """
    Genera un DOCX con título y párrafos de texto.

    Args:
        path: Ruta de salida.
        titulo: Título del documento.
        contenido: Texto del cuerpo.

    Returns:
        Ruta del archivo generado.
    """
    try:
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc = Document()
        doc.add_heading(titulo, level=0)
        for parrafo in contenido.split("\n\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())
        doc.save(path)
        logger.info("DOCX generado: %s", path)
        return path
    except Exception as e:
        logger.error("Error generando Word: %s", e)
        raise


def generar_word_desde_plantilla(plantilla_path: str, path_salida: str, variables: dict) -> str:
    """
    Genera DOCX reemplazando placeholders {{CLAVE}} en plantilla de texto.

    Args:
        plantilla_path: Ruta de plantilla .txt.
        path_salida: Ruta del DOCX de salida.
        variables: Dict de reemplazos.

    Returns:
        Ruta del archivo generado.
    """
    try:
        with open(plantilla_path, encoding="utf-8") as f:
            texto = f.read()
        for clave, valor in variables.items():
            texto = texto.replace("{{" + clave.upper() + "}}", str(valor))
        if "{{TEXTO_CLAUDE}}" not in plantilla_path and "TEXTO_CLAUDE" in variables:
            texto = variables["TEXTO_CLAUDE"]
        return generar_word_texto(path_salida, variables.get("NOMBRE_EMPRESA", "Documento"), texto)
    except Exception as e:
        logger.error("Error con plantilla Word: %s", e)
        return generar_word_texto(path_salida, "Documento", str(variables.get("TEXTO_CLAUDE", "")))
